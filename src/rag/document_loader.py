"""
Document Loader
多格式文档解析器 — 支持 PDF、Word、TXT 批量导入与自动解析清洗
"""

import re
import asyncio
from pathlib import Path
from typing import List, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime


@dataclass
class Document:
    """文档数据结构"""
    id: str
    filename: str
    file_type: str  # pdf / word / txt
    content: str
    chunks: List[str] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    created_at: str = field(default_factory=lambda: datetime.now().isoformat())

    def to_dict(self) -> dict:
        return {
            'id': self.id,
            'filename': self.filename,
            'file_type': self.file_type,
            'content_preview': self.content[:200],
            'chunk_count': len(self.chunks),
            'metadata': self.metadata,
            'created_at': self.created_at
        }


class DocumentLoader:
    """多格式文档加载与解析器"""

    def __init__(self, config: dict = None):
        self.config = config or {}
        self.chunk_size = self.config.get('chunk_size', 500)
        self.chunk_overlap = self.config.get('chunk_overlap', 50)

    async def load_file(self, file_path: str) -> Optional[Document]:
        """加载单个文件并自动识别格式"""
        path = Path(file_path)
        if not path.exists():
            print(f"[DocLoader] File not found: {file_path}")
            return None

        suffix = path.suffix.lower()
        if suffix == '.pdf':
            return await self._load_pdf(path)
        elif suffix in ('.docx', '.doc'):
            return await self._load_word(path)
        elif suffix in ('.txt', '.md', '.py', '.js', '.yaml', '.json', '.csv'):
            return await self._load_text(path)
        else:
            print(f"[DocLoader] Unsupported format: {suffix}")
            return None

    async def load_directory(self, dir_path: str) -> List[Document]:
        """批量加载目录下所有支持的文档"""
        path = Path(dir_path)
        if not path.is_dir():
            return []

        tasks = []
        for f in path.rglob('*'):
            if f.suffix.lower() in ('.pdf', '.docx', '.doc', '.txt', '.md'):
                tasks.append(self.load_file(str(f)))

        results = await asyncio.gather(*tasks)
        return [doc for doc in results if doc is not None]

    async def _load_pdf(self, path: Path) -> Document:
        """解析PDF文档 — Docling优先，PyMuPDF兜底"""
        content, tables_data, page_images = "", [], []

        # 尝试 Docling（IBM开源，MIT协议，结构化提取最优）
        try:
            from docling.document_converter import DocumentConverter
            loop = asyncio.get_event_loop()

            def extract_docling():
                converter = DocumentConverter()
                result = converter.convert(str(path))
                docling_doc = result.document
                # 导出为结构化markdown（含表格）
                md = docling_doc.export_to_markdown()
                # 提取表格
                _tables = []
                for table in docling_doc.tables:
                    if table.data:
                        _tables.append({
                            'caption': table.caption_text if hasattr(table, 'caption_text') else '',
                            'rows': len(table.data.rows) if table.data.rows else 0
                        })
                return md, _tables
            content, tables_data = await loop.run_in_executor(None, extract_docling)
            if content:
                print(f"[DocLoader] Docling parsed: {path.name} ({len(content)} chars)")

                # 渲染页面截图（用于视觉检索）
                try:
                    page_images = await self._render_page_images(path)
                except Exception as e:
                    print(f"[DocLoader] Page rendering failed: {e}")

                doc = self._build_document(path, content, 'pdf')
                doc.metadata['parser'] = 'docling'
                doc.metadata['tables_count'] = len(tables_data)
                doc.metadata['page_images'] = page_images
                doc.metadata['has_visual'] = len(page_images) > 0
                return doc
        except ImportError:
            print(f"[DocLoader] Docling not installed, falling back to PyMuPDF")
        except Exception as e:
            print(f"[DocLoader] Docling failed: {e}, falling back to PyMuPDF")

        # Fallback: PyMuPDF
        return await self._load_pdf_fitz(path)

    async def _load_pdf_fitz(self, path: Path) -> Document:
        """PyMuPDF兜底解析PDF"""
        content, tables_data = "", []
        try:
            import fitz
            loop = asyncio.get_event_loop()

            def extract():
                text_parts = []
                doc = fitz.open(str(path))
                meta = doc.metadata
                if meta.get('title'):
                    text_parts.append(f"# {meta['title']}")
                for page_num, page in enumerate(doc, 1):
                    page_text = page.get_text('text')
                    text_parts.append(page_text)
                    try:
                        tables = page.find_tables()
                        if tables:
                            for t_idx, table in enumerate(tables, 1):
                                rows = table.extract()
                                if rows:
                                    text_parts.append(f"\n[表格 {page_num}-{t_idx}]")
                                    headers = [str(h or '') for h in rows[0]]
                                    text_parts.append('| ' + ' | '.join(headers) + ' |')
                                    text_parts.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
                                    for row in rows[1:]:
                                        cells = [str(c or '') for c in row]
                                        text_parts.append('| ' + ' | '.join(cells) + ' |')
                                    tables_data.append({
                                        'page': page_num,
                                        'table_index': t_idx
                                    })
                    except Exception:
                        pass
                page_images = self._render_page_images_fitz(doc)
                doc.close()
                return '\n'.join(text_parts), tables_data, page_images

            content, tables_data, page_images = await loop.run_in_executor(None, extract)
        except ImportError:
            content = f"[PDF解析不可用，请安装docling或pymupdf: {path.name}]"
        except Exception as e:
            print(f"[DocLoader] PyMuPDF error: {e}")
            content = f"[PDF解析错误: {path.name}]"

        doc = self._build_document(path, content, 'pdf')
        doc.metadata['parser'] = 'pymupdf'
        doc.metadata['tables_count'] = len(tables_data)
        doc.metadata['page_images'] = page_images
        doc.metadata['has_visual'] = len(page_images) > 0
        return doc

    def _render_page_images_fitz(self, doc) -> List[str]:
        """PyMuPDF渲染页面为base64图片"""
        images = []
        try:
            for page_num in range(min(len(doc), 20)):  # 最多20页
                page = doc[page_num]
                pix = page.get_pixmap(dpi=150)
                import base64
                img_b64 = base64.b64encode(pix.tobytes('png')).decode()
                images.append(f"data:image/png;base64,{img_b64}")
        except Exception:
            pass
        return images

    async def _render_page_images(self, path: Path) -> List[str]:
        """渲染PDF页面为base64图片（用于视觉检索）"""
        import base64
        try:
            import fitz
            loop = asyncio.get_event_loop()

            def render():
                imgs = []
                doc = fitz.open(str(path))
                for page_num in range(min(len(doc), 20)):
                    page = doc[page_num]
                    pix = page.get_pixmap(dpi=150)
                    imgs.append(base64.b64encode(pix.tobytes('png')).decode())
                doc.close()
                return imgs

            return await loop.run_in_executor(None, render)
        except ImportError:
            return []

    async def _load_word(self, path: Path) -> Document:
        """解析Word文档"""
        content = ""
        try:
            from docx import Document as DocxDocument
            loop = asyncio.get_event_loop()

            def extract():
                doc = DocxDocument(str(path))
                return '\n'.join(p.text for p in doc.paragraphs)

            content = await loop.run_in_executor(None, extract)
        except ImportError:
            content = f"[Word content placeholder: {path.name}]"
        except Exception as e:
            print(f"[DocLoader] Word parse error: {e}")
            content = f"[Word parse error for {path.name}]"

        return self._build_document(path, content, 'word')

    async def _load_text(self, path: Path) -> Document:
        """解析纯文本/TXT/Markdown等"""
        content = ""
        try:
            loop = asyncio.get_event_loop()

            def read_file():
                return path.read_text(encoding='utf-8', errors='replace')

            content = await loop.run_in_executor(None, read_file)
        except Exception as e:
            print(f"[DocLoader] Text read error: {e}")
            content = f"[Read error for {path.name}]"

        return self._build_document(path, content, path.suffix.lstrip('.'))

    def _build_document(self, path: Path, raw_content: str, file_type: str) -> Document:
        """构建Document对象并进行清洗分块"""
        cleaned = self._clean_text(raw_content)
        chunks = self._chunk_text(cleaned)

        import uuid
        return Document(
            id=str(uuid.uuid4()),
            filename=path.name,
            file_type=file_type,
            content=cleaned,
            chunks=chunks,
            metadata={
                'source_path': str(path),
                'file_size': path.stat().st_size if path.exists() else 0,
                'char_count': len(cleaned),
                'chunk_count': len(chunks)
            }
        )

    def _clean_text(self, text: str) -> str:
        """文本清洗"""
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'\r', '\n', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        text = re.sub(r'\t+', ' ', text)
        # 移除控制字符（保留常见换行制表符）
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()

    def _chunk_text(self, text: str) -> List[str]:
        """语义分块：按段落→句子边界切分，保持语义完整"""
        if len(text) <= self.chunk_size:
            return [text] if text.strip() else []

        # 1. 提取文档层级结构（标题/章节）
        sections = self._split_by_sections(text)

        # 2. 对每个 section 进行段落→句子粒度分块
        chunks = []
        for section_title, section_text in sections:
            section_chunks = self._semantic_chunk(section_text, section_title)
            chunks.extend(section_chunks)

        return chunks

    def _split_by_sections(self, text: str) -> List[Tuple[str, str]]:
        """按标题/章节分割文档，保留层级上下文"""
        # 匹配 Markdown 标题、中文章节标记等
        section_patterns = [
            r'^#{1,3}\s+(.+)$',           # Markdown headers
            r'^(第[一二三四五六七八九十\d]+[章节篇部])[：:\s]*(.*)$',  # 中文章节
            r'^(\d+[\.、]\s*.+)$',          # 数字标题
        ]

        lines = text.split('\n')
        sections = []
        current_title = ''
        current_lines = []

        for line in lines:
            is_section = False
            for pattern in section_patterns:
                m = re.match(pattern, line.strip())
                if m:
                    if current_lines:
                        sections.append((current_title, '\n'.join(current_lines)))
                    current_title = line.strip()
                    current_lines = []
                    is_section = True
                    break
            if not is_section:
                current_lines.append(line)

        if current_lines:
            sections.append((current_title, '\n'.join(current_lines)))

        if not sections:
            sections = [('', text)]

        return sections

    def _semantic_chunk(self, text: str, section_title: str = '') -> List[str]:
        """按段落→句子语义边界分块"""
        # 按段落分割
        paragraphs = re.split(r'\n\s*\n', text)
        # 将每个段落拆成句子
        sentences = []
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            para_sents = self._split_sentences(para)
            sentences.extend(para_sents)

        if not sentences:
            return []

        chunks = []
        current_chunk = ''
        prefix = f"[{section_title}] " if section_title else ''

        for sent in sentences:
            # 单句加入后是否超限
            if len(current_chunk) + len(sent) <= self.chunk_size:
                current_chunk += sent
            else:
                # 保存当前块
                if current_chunk.strip():
                    chunks.append(prefix + current_chunk.strip())
                # 新块开始，带重叠
                if current_chunk and self.chunk_overlap > 0:
                    overlap_text = current_chunk[-self.chunk_overlap:]
                    current_chunk = overlap_text + sent
                else:
                    current_chunk = sent

                # 长句单独成块
                while len(current_chunk) > self.chunk_size * 1.5:
                    split_at = current_chunk.rfind('，', 0, self.chunk_size)
                    if split_at < self.chunk_size // 2:
                        split_at = current_chunk.rfind('；', 0, self.chunk_size)
                    if split_at < self.chunk_size // 2:
                        split_at = self.chunk_size
                    chunks.append(prefix + current_chunk[:split_at].strip())
                    current_chunk = current_chunk[max(0, split_at - self.chunk_overlap):]

        if current_chunk.strip():
            chunks.append(prefix + current_chunk.strip())

        return chunks

    def _split_sentences(self, text: str) -> List[str]:
        """句子边界切分（中英文）"""
        # 在句末标点后切分
        pattern = r'(?<=[。！？.!?\n])(?=[^\s])'
        parts = re.split(pattern, text)
        # 确保标点跟在句子末尾
        result = []
        buffer = ''
        for part in parts:
            buffer += part
            if re.search(r'[。！？.!?]$', buffer.strip()):
                result.append(buffer)
                buffer = ''
        if buffer.strip():
            result.append(buffer)
        return result if result else [text]

    async def load_file_with_ocr(self, file_path: str, ocr_lang: str = 'chi_sim+eng') -> Optional[Document]:
        """加载图片文件并进行OCR"""
        path = Path(file_path)
        if not path.exists():
            return None

        suffix = path.suffix.lower()
        if suffix in ('.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.webp'):
            return await self._ocr_image(path, ocr_lang)
        else:
            return await self.load_file(file_path)

    async def _ocr_image(self, path: Path, lang: str) -> Optional[Document]:
        """OCR图片提取文字"""
        content = ""
        try:
            # 尝试使用 paddleocr / easyocr / tesseract
            try:
                from paddleocr import PaddleOCR
                loop = asyncio.get_event_loop()

                def paddle_ocr():
                    ocr = PaddleOCR(lang=lang, use_angle_cls=True)
                    result = ocr.ocr(str(path))
                    texts = []
                    if result and result[0]:
                        for line in result[0]:
                            texts.append(line[1][0])
                    return '\n'.join(texts)

                content = await loop.run_in_executor(None, paddle_ocr)
            except ImportError:
                try:
                    import pytesseract
                    from PIL import Image
                    loop = asyncio.get_event_loop()

                    def tesseract_ocr():
                        img = Image.open(str(path))
                        return pytesseract.image_to_string(img, lang='chi_sim+eng')

                    content = await loop.run_in_executor(None, tesseract_ocr)
                except ImportError:
                    content = f"[OCR未就绪，图片内容占位: {path.name}]"
        except Exception as e:
            print(f"[DocLoader] OCR error: {e}")
            content = f"[OCR错误: {path.name}]"

        return self._build_document(path, content, path.suffix.lstrip('.'))

    def get_chunk_statistics(self, document: Document) -> dict:
        """获取文档分块统计"""
        return {
            'filename': document.filename,
            'total_chars': len(document.content),
            'chunk_count': len(document.chunks),
            'avg_chunk_size': sum(len(c) for c in document.chunks) / len(document.chunks) if document.chunks else 0,
            'file_type': document.file_type
        }
